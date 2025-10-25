from django.shortcuts import render, redirect, get_object_or_404
from django.http import HttpResponse
from django.template.loader import get_template
from django.utils import timezone
from .models import Invoice, LetterHead
from .forms import InvoiceForm, InvoiceItemFormSet, LetterHeadForm
# reportlab.pdfgen import moved into generate_pdf to avoid import-time errors if reportlab is not installed
from reportlab.lib.pagesizes import letter
from docx import Document
import io

def number_to_words(number):
    def get_words(n):
        units = ["", "One", "Two", "Three", "Four", "Five", "Six", "Seven", "Eight", "Nine", "Ten", "Eleven", "Twelve", "Thirteen", "Fourteen", "Fifteen", "Sixteen", "Seventeen", "Eighteen", "Nineteen"]
        tens = ["", "", "Twenty", "Thirty", "Forty", "Fifty", "Sixty", "Seventy", "Eighty", "Ninety"]
        if n == 0: return ""
        if n < 20: return units[n]
        if n < 100: return tens[n // 10] + (" " + units[n % 10] if n % 10 != 0 else "")
        if n < 1000: return units[n // 100] + " Hundred" + (" and " + get_words(n % 100) if n % 100 != 0 else "")
        if n < 100000: return get_words(n // 1000) + " Thousand" + (" " + get_words(n % 1000) if n % 1000 != 0 else "")
        return get_words(n // 100000) + " Lakh" + (" " + get_words(n % 100000) if n % 100000 != 0 else "")
    
    return get_words(int(number))

def letterhead_list(request):
    if request.method == 'POST':
        form = LetterHeadForm(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            return redirect('letterhead_list')
    else:
        form = LetterHeadForm()
    
    letterheads = LetterHead.objects.all()
    return render(request, 'invoice_app/letterhead_list.html', {
        'form': form,
        'letterheads': letterheads
    })

def create_invoice(request):
    if request.method == 'POST':
        form = InvoiceForm(request.POST)
        if form.is_valid():
            invoice = form.save()
            formset = InvoiceItemFormSet(request.POST, instance=invoice)
            if formset.is_valid():
                formset.save()
                return redirect('invoice_detail', pk=invoice.pk)
    else:
        form = InvoiceForm()
        formset = InvoiceItemFormSet()
    
    return render(request, 'invoice_app/create_invoice.html', {
        'form': form,
        'formset': formset
    })

def invoice_detail(request, pk):
    invoice = get_object_or_404(Invoice, pk=pk)
    return render(request, 'invoice_app/invoice_detail.html', {
        'invoice': invoice
    })

def generate_pdf(request, pk):
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import inch, mm
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import Paragraph, Table, TableStyle
        from reportlab.lib import colors
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except ImportError:
        return HttpResponse("reportlab is not installed; install it with: pip install reportlab", status=500)

    invoice = get_object_or_404(Invoice, pk=pk)
    buffer = io.BytesIO()
    
    # Create the PDF object using A4 size
    width, height = A4
    p = canvas.Canvas(buffer, pagesize=A4)
    
    margin_left = 40
    margin_right = width - 40
    
    # Add letterhead
    if invoice.letterhead.image:
        # Get the image size
        img_width = width - 2*margin_left  # Full width minus margins
        img_height = 100  # Fixed height for letterhead
        p.drawImage(invoice.letterhead.image.path, margin_left, height - img_height, width=img_width, height=img_height, preserveAspectRatio=True, mask='auto')
    
    # Add document title centered
    doc_title = "Invoice" if invoice.document_type == "invoice" else "Quotation"
    p.setFont("Helvetica-Bold", 11)
    p.drawString(width/2 - 30, height - 130, doc_title)
    
    # Two-column layout for To and Date
    p.setFont("Helvetica-Bold", 10)
    p.drawString(margin_left, height - 160, "TO")
    p.drawString(margin_right - 150, height - 160, f"Date: {invoice.date.strftime('%d/%m/%Y')}")
    
    # Customer details
    y = height - 180  # Start position for customer details
    
    # Customer name in bold
    p.setFont("Helvetica-Bold", 10)
    for line in invoice.customer_name.split('\n'):
        p.drawString(margin_left + 20, y, line.strip())
        y -= 15
    
    # Address with proper indentation
    p.setFont("Helvetica", 10)
    for line in invoice.address.split('\n'):
        p.drawString(margin_left + 20, y, line.strip())
        y -= 15
    
    y -= 10  # Reduced space before table
    
    # Create table data
    data = [['Sr.\nNo.', 'Area\nSite', 'Length', 'Breadth', 'Quantity', 'Area', 'Rate per\nsq.ft.', 'Amount']]
    
    total = 0
    for item in invoice.items.all():
        total += item.total
        data.append([
            str(item.srno),
            item.description,
            f"{item.length}",
            f"{item.breadth}",
            f"{item.quantity}",
            f"{item.area}",
            f"{item.unit_price} Rs",
            f"{item.total} Rs"
        ])
    
    # Add total amount row
    data.append(['', '', '', '', '', '', 'All Total Amount', f"{total} Rs"])
    
    # Add Grand Total row with text in words
    data.append(['', '', '', '', '', '', 'Grand Total', f"{total} Rs"])
    data.append(['in words', f"Rupees {number_to_words(int(total))} Only", '', '', '', '', '', ''])
    
    # Calculate column widths to match the image
    table = Table(data, colWidths=[35, 80, 45, 45, 45, 45, 70, 70])
    table.setStyle(TableStyle([
        ('FONT', (0, 0), (-1, -1), 'Helvetica', 9),
        ('FONT', (0, 0), (-1, 0), 'Helvetica-Bold', 9),  # Header row
        ('FONT', (-2, -3), (-1, -2), 'Helvetica-Bold', 9),  # Total rows
        ('ALIGN', (0, 0), (0, -1), 'CENTER'),  # Sr.No. centered
        ('ALIGN', (2, 1), (-1, -2), 'CENTER'),  # Numbers centered
        ('GRID', (0, 0), (-1, -4), 1, colors.black),  # Grid lines except for total rows
        ('LINEABOVE', (0, -3), (-1, -3), 1, colors.black),  # Line above total
        ('LINEBELOW', (0, -2), (-1, -2), 1, colors.black),  # Line below grand total
        ('SPAN', (1, -1), (-1, -1)),  # Span the 'in words' row
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
    ]))
    
    # Draw table with proper spacing
    table.wrapOn(p, width, height)
    table_height = len(data) * 20  # Approximate height based on number of rows
    table.drawOn(p, 50, y - table_height - 30)  # Adjust position based on table size
    
    # Calculate remaining space
    remaining_height = y - table_height - 150  # Space between table and bottom margin
    
    # Position note and signature with proper spacing
    note_y = remaining_height - 30  # Start note section higher up
    
    p.setFont("Helvetica-Bold", 10)
    p.drawString(margin_left, note_y, "NOTE:")
    p.setFont("Helvetica", 9)
    
    # Get notes from the invoice if available, otherwise use defaults
    notes = invoice.note.split('\n') if invoice.note else [
        "1. Payment: 50% advance with order, balance before delivery",
        "2. Delivery Period: 4-5 weeks from the date of order confirmation",
        "3. GST Extra as applicable",
        "4. Validity: This quotation is valid for 30 days",
        "5. Installation charges extra if required"
    ]
    
    # Draw notes with proper spacing
    for note in notes:
        if note.strip():
            note_y -= 15
            p.drawString(margin_left, note_y, note.strip())
    
    # Add footer (signature) on the right side with balanced spacing
    sig_y = note_y - 40  # Space between notes and signature
    p.setFont("Helvetica", 10)
    p.drawString(margin_right - 168, sig_y, "From " + invoice.letterhead.name)
    p.drawString(margin_right - 170 , sig_y - 20, "Authorized Signatory")
    
    # Save the PDF
    p.showPage()
    p.save()
    
    buffer.seek(0)
    response = HttpResponse(buffer, content_type='application/pdf')
    filename_prefix = "Invoice" if invoice.document_type == "invoice" else "Quotation"
    response['Content-Disposition'] = f'attachment; filename={filename_prefix}_{invoice.customer_name}.pdf'
    return response

def generate_word(request, pk):
    invoice = get_object_or_404(Invoice, pk=pk)
    doc = Document()
    
    # Add letterhead info
    doc.add_heading(invoice.letterhead.name, 0)
    
    # Add invoice details
    doc.add_paragraph(f"To: {invoice.customer_name}")
    doc.add_paragraph(invoice.address)
    doc.add_paragraph(f"Date: {invoice.date}")
    
    # Add items table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Description'
    header_cells[1].text = 'Quantity'
    header_cells[2].text = 'Unit Price'
    header_cells[3].text = 'Total'
    
    for item in invoice.items.all():
        row_cells = table.add_row().cells
        row_cells[0].text = item.description
        row_cells[1].text = str(item.quantity)
        row_cells[2].text = str(item.unit_price)
        row_cells[3].text = str(item.total)
    
    # Calculate and add total
    total = sum(item.total for item in invoice.items.all())
    doc.add_paragraph(f"Grand Total: {total}")
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(buffer.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename=invoice_{pk}.docx'
    return response